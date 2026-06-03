"""
backend/processor.py

OOP 封装视频处理与 Word 文档生成逻辑。

Classes:
    ProcessParams  — 处理参数 dataclass
    VideoProcessor — 视频帧提取 + 哈希去重
    DocxBuilder    — Word 文档生成
"""

from __future__ import annotations

import asyncio
import json
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import cv2
import imagehash
import numpy as np
from PIL import Image

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


# ──────────────────────────────────────────────────────────────────────────────
# 数据结构
# ──────────────────────────────────────────────────────────────────────────────

@dataclass
class ProcessParams:
    """视频处理的可调参数。"""
    sample_fps: int = 5          # 每秒采样帧数
    change_threshold: float = 3  # 帧差阈值（像素均值）
    stable_seconds: float = 2    # 连续稳定秒数
    hash_threshold: int = 5      # 感知哈希距离阈值


@dataclass
class ProgressEvent:
    """SSE 进度事件，可序列化为 JSON。"""
    type: str                          # "progress" | "done" | "error"
    message: str = ""
    current: int = 0
    total: int = 0
    saved: int = 0

    def to_json(self) -> str:
        return json.dumps(self.__dict__, ensure_ascii=False)


# ──────────────────────────────────────────────────────────────────────────────
# 视频处理器
# ──────────────────────────────────────────────────────────────────────────────

class VideoProcessor:
    """
    从视频中提取稳定、不重复的帧并保存为图片。

    Usage:
        processor = VideoProcessor(video_path, output_dir, params)
        await processor.run(queue)
    """

    def __init__(
        self,
        video_path: Path,
        output_dir: Path,
        params: ProcessParams,
    ) -> None:
        self._video_path = video_path
        self._output_dir = output_dir
        self._params = params
        self._output_dir.mkdir(parents=True, exist_ok=True)

    # ── 公共接口 ──────────────────────────────────────

    async def run(self, queue: asyncio.Queue[ProgressEvent]) -> int:
        """
        异步执行视频分析，通过 queue 推送 ProgressEvent。

        Returns:
            保存的截图数量。

        Raises:
            RuntimeError: 视频文件无法读取时。
        """
        loop = asyncio.get_event_loop()
        try:
            saved = await loop.run_in_executor(None, self._process_sync, queue, loop)
        except Exception as exc:
            await queue.put(ProgressEvent(type="error", message=str(exc)))
            raise
        return saved

    # ── 内部实现 ──────────────────────────────────────

    def _process_sync(
        self,
        queue: asyncio.Queue[ProgressEvent],
        loop: asyncio.AbstractEventLoop,
    ) -> int:
        """在线程执行器中同步运行的帧分析逻辑。"""
        cap = cv2.VideoCapture(str(self._video_path))
        if not cap.isOpened():
            raise RuntimeError(f"无法打开视频：{self._video_path}")

        fps: float = cap.get(cv2.CAP_PROP_FPS) or 25.0
        total_frames: int = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        frame_interval: int = max(1, int(fps / self._params.sample_fps))
        stable_needed: int = int(self._params.sample_fps * self._params.stable_seconds)

        ret, prev_frame = cap.read()
        if not ret:
            cap.release()
            raise RuntimeError("视频第一帧读取失败")

        stable_count = 0
        saved_count = 0
        last_hash: imagehash.ImageHash | None = None
        frame_index = 0

        def _push(event: ProgressEvent) -> None:
            asyncio.run_coroutine_threadsafe(queue.put(event), loop)

        _push(ProgressEvent(type="progress", message="开始分析视频…", total=total_frames))

        while True:
            ret, frame = cap.read()
            if not ret:
                break

            frame_index += 1

            if frame_index % frame_interval != 0:
                continue

            stable_count, saved_count, last_hash = self._analyze_frame(
                prev_frame, frame, stable_count, stable_needed,
                saved_count, last_hash, _push, total_frames, frame_index,
            )

            prev_frame = frame.copy()

        cap.release()
        _push(ProgressEvent(type="done", saved=saved_count, message=f"提取完成，共 {saved_count} 张截图"))
        return saved_count

    def _analyze_frame(
        self,
        prev_frame: np.ndarray,
        frame: np.ndarray,
        stable_count: int,
        stable_needed: int,
        saved_count: int,
        last_hash: imagehash.ImageHash | None,
        push,
        total_frames: int,
        frame_index: int,
    ) -> tuple[int, int, imagehash.ImageHash | None]:
        """分析单帧，必要时保存。返回 (stable_count, saved_count, last_hash)。"""
        gray1 = cv2.cvtColor(prev_frame, cv2.COLOR_BGR2GRAY)
        gray2 = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        score: float = float(np.mean(cv2.absdiff(gray1, gray2)))

        if score < self._params.change_threshold:
            stable_count += 1
        else:
            stable_count = 0

        push(ProgressEvent(
            type="progress",
            current=frame_index,
            total=total_frames,
            saved=saved_count,
            message=f"分析第 {frame_index} 帧",
        ))

        if stable_count >= stable_needed:
            saved_count, last_hash = self._try_save_frame(frame, saved_count, last_hash)
            stable_count = 0

        return stable_count, saved_count, last_hash

    def _try_save_frame(
        self,
        frame: np.ndarray,
        saved_count: int,
        last_hash: imagehash.ImageHash | None,
    ) -> tuple[int, imagehash.ImageHash | None]:
        """计算感知哈希，若与上一张差异足够大则保存。"""
        temp_file = self._output_dir / f"temp_{saved_count}.png"
        cv2.imwrite(str(temp_file), frame)

        try:
            img = Image.open(temp_file)
            current_hash = imagehash.phash(img)
        finally:
            img.close()

        should_save = (last_hash is None) or (
            (current_hash - last_hash) > self._params.hash_threshold
        )

        if should_save:
            final_file = self._output_dir / f"page_{saved_count:03d}.png"
            os.rename(temp_file, final_file)
            return saved_count + 1, current_hash
        else:
            os.remove(temp_file)
            return saved_count, last_hash


# ──────────────────────────────────────────────────────────────────────────────
# Word 文档构建器
# ──────────────────────────────────────────────────────────────────────────────

class DocxBuilder:
    """
    将截图目录下的图片整理成 Word 文档（每行 cols 张，每页 cols×2 张）。

    Usage:
        builder = DocxBuilder(screenshots_dir, output_path, cols=2)
        builder.build()
    """

    ROWS_PER_PAGE: int = 2

    def __init__(self, screenshots_dir: Path, output_path: Path, cols: int = 2) -> None:
        self._screenshots_dir = screenshots_dir
        self._output_path = output_path
        self.COLS = cols
        self.IMAGE_WIDTH_CM = round(14.0 / cols, 1)

    # ── 公共接口 ──────────────────────────────────────

    def build(self) -> Path:
        """构建 Word 文档并写入磁盘，返回输出路径。"""
        images = sorted(self._screenshots_dir.glob("page_*.png"))
        if not images:
            raise RuntimeError("截图目录中没有找到图片")

        doc = Document()
        doc.add_heading("视频截图整理", level=1)

        images_per_page = self.COLS * self.ROWS_PER_PAGE

        for page_idx in range(0, len(images), images_per_page):
            if page_idx > 0:
                doc.add_page_break()
            page_images = images[page_idx: page_idx + images_per_page]
            self._add_image_table(doc, page_images)

        doc.save(str(self._output_path))
        return self._output_path

    # ── 内部实现 ──────────────────────────────────────

    def _add_image_table(self, doc: Document, page_images: list[Path]) -> None:
        """在文档中插入一个无边框图片表格。"""
        rows_needed = (len(page_images) + self.COLS - 1) // self.COLS
        table = doc.add_table(rows=rows_needed, cols=self.COLS)
        self._remove_table_borders(table)

        for idx, img_path in enumerate(page_images):
            row_idx, col_idx = divmod(idx, self.COLS)
            cell = table.cell(row_idx, col_idx)
            self._set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            self._zero_paragraph_spacing(paragraph)
            paragraph.add_run().add_picture(str(img_path), width=Cm(self.IMAGE_WIDTH_CM))

    @staticmethod
    def _remove_table_borders(table: Any) -> None:
        tbl = table._tbl
        tblPr = tbl.tblPr
        tbl_borders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "none")
            tbl_borders.append(border)
        tblPr.append(tbl_borders)

    @staticmethod
    def _set_cell_margins(
        cell: Any,
        top: int = 0,
        left: int = 30,
        bottom: int = 0,
        right: int = 30,
    ) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_mar = OxmlElement("w:tcMar")
        for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
            node = OxmlElement(f"w:{side}")
            node.set(qn("w:w"), str(val))
            node.set(qn("w:type"), "dxa")
            tc_mar.append(node)
        tc_pr.append(tc_mar)

    @staticmethod
    def _zero_paragraph_spacing(paragraph: Any) -> None:
        pf = paragraph.paragraph_format
        pf.space_before = Cm(0)
        pf.space_after = Cm(0)
